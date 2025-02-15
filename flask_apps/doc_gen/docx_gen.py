# docx_gen.py
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree
import re
import base64
from io import BytesIO
import markdown
from bs4 import BeautifulSoup, Tag
from docx.shared import Inches, Length  # Import Inches and Length
from docx.oxml.table import CT_TblPr


# --- Configuration ---
PAGE_MARGIN = Cm(2)
# Calculate editable_width and explicitly convert to Inches
editable_width_cm_val = Cm(21).cm - (PAGE_MARGIN.cm + PAGE_MARGIN.cm)
editable_width = Cm(editable_width_cm_val)  # Ensure editable_width is Inches type


# Load XSLT for MathML to OMML conversion
MML2OMML_STYLESHEET_PATH = 'MML2OMML.XSL'  # Ensure this file is in the same directory
mml2omml_stylesheet = etree.parse(MML2OMML_STYLESHEET_PATH)
mml2omml_transform = etree.XSLT(mml2omml_stylesheet)


def add_markdown_table(doc, table_element):
    """Convert BeautifulSoup table element to Word table."""
    rows = table_element.find_all('tr')
    cols = max(len(row.find_all(['th', 'td'])) for row in rows)
    word_table = doc.add_table(rows=0, cols=cols)
    word_table.style = 'Table Grid'

    for html_row in rows:
        cells = html_row.find_all(['th', 'td'])
        word_row = word_table.add_row()
        for idx, cell in enumerate(cells):
            add_markdown_to_doc(doc, cell.decode_contents(), parent=word_row.cells[idx])
    return word_table

def add_markdown_to_doc(doc, markdown_text, parent=None):
    """Add Markdown content with robust formatting."""
    html = markdown.markdown(markdown_text, extensions=['tables', 'extra'])
    soup = BeautifulSoup(html, 'html.parser')

    container = parent or doc
    p = container.add_paragraph()
    p.paragraph_format.keep_together = True  # Add keep_together here

    def process_element(element, current_para=None, format_name=None):
        current_para = current_para or p
        if isinstance(element, str):
            handle_text(element, current_para, format_name)
            return

        if element.name in ['strong', 'em', 'del', 'u', 'code']:
            run = current_para.add_run()
            apply_formatting(run, element.name)
            for child in element.contents:
                process_element(child, current_para, element.name)
        elif element.name == 'table':
            add_markdown_table(doc, element)
        # No special handling for other elements, process children directly
        else:
            for child in element.contents:
                process_element(child, current_para)


    def handle_text(text, parent_para, format_name):
        segments = re.split(r'(\\\(.*?\\\)|\\\[.*?\\\]|\$.*?\$)', text)
        for seg in segments:
            if not seg:  # Skip empty segments
                continue

            if re.match(r'(\\\(.*?\\\)|\\\[.*?\\\]|\$.*?\$)', seg):
                # It's a LaTeX segment
                latex_content = re.sub(r'^\$|\\\(|\\\[|\$|\\\)|\\\]$', '', seg)
                try:
                    mathml = latex_to_mathml(latex_content)
                    tree = etree.fromstring(mathml)
                    omml = mml2omml_transform(tree)
                    run = parent_para.add_run()
                    run._element.append(parse_xml(etree.tostring(omml)))

                except Exception as e:
                    # Handle LaTeX conversion errors gracefully
                    error_run = parent_para.add_run(f'[Math Error: {latex_content}]')
                    error_run.font.color.rgb = (0xFF, 0x00, 0x00)  # Red color for errors

            else:
                # Regular text segment
                run = parent_para.add_run(seg)
                if format_name:
                    apply_formatting(run, format_name)

    def apply_formatting(run, tag):
        formats = {
            'strong': {'bold': True},
            'em': {'italic': True},
            'del': {'strike': True},
            'u': {'underline': True},
            'code': {'font': 'Courier New'}  # Example of setting a specific font
        }
        for prop, val in formats.get(tag, {}).items():
            setattr(run, prop, val) if prop != 'font' else setattr(run.font, 'name', val)

    process_element(soup)
    return p

def insertHR(paragraph, left_margin=Cm(1)):
    """
    Inserts a horizontal rule below the given paragraph with an optional left margin.
    """
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()

    # --- Add left margin (indentation) ---
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(left_margin.twips)))  # Set left indent in twips
    pPr.insert_element_before(ind,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )

    # --- Add bottom border (horizontal rule) ---
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:ind', # Insert pBdr before 'w:ind' now that we've added it above
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def remove_table_padding(table):
    """
    Removes padding from all cells in a python-docx table.
    """
    for row in table.rows:
        for cell in row.cells:
            _remove_cell_padding(cell._tc)  # Access the underlying <w:tc> element

def _remove_cell_padding(tc):
    """Removes padding from a single table cell's <w:tc> element."""
    tcPr = tc.get_or_add_tcPr() # Get or create tcPr (table cell properties) element
    tcMar = OxmlElement('w:tcMar') # Create tcMar (table cell margins) element

    # Set all margins to zero
    top = OxmlElement('w:top')
    top.set(qn('w:w'), '0')  # '0' twips
    top.set(qn('w:type'), 'dxa') # Required attribute

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:w'), '0')
    bottom.set(qn('w:type'), 'dxa')

    left = OxmlElement('w:left')
    left.set(qn('w:w'), '0')
    left.set(qn('w:type'), 'dxa')

    right = OxmlElement('w:right')
    right.set(qn('w:w'), '0')
    right.set(qn('w:type'), 'dxa')

    # Append margin elements to tcMar
    tcMar.append(top)
    tcMar.append(bottom)
    tcMar.append(left)
    tcMar.append(right)

    # Insert tcMar into tcPr, overwriting any existing tcMar
    tcPr.replace(tcMar) # Or tcPr.append(tcMar) if you want to add instead of replace

def set_table_indentation(table, left_indent):
    """
    Sets the left indentation of a python-docx table.
    """
    if left_indent < Length(0):  # Ensure indent is not negative
        raise ValueError("Left indent must be a non-negative Length object.")

    table_xml = table._tbl  # Get the underlying <w:tbl> XML element

    # --- Correct way to get or add <w:tblPr> ---
    table_properties = table_xml.tblPr  # Try to get existing <w:tblPr>
    if table_properties is None:
        table_properties = CT_TblPr.new_tblPr() # Create a new <w:tblPr>
        table_xml.insert_element_before(table_properties, 'tblGrid') # Insert it into <w:tbl>

    table_indentation = OxmlElement('w:tblInd') # Create <w:tblInd> element
    table_indentation.set(qn('w:w'), str(int(left_indent.twips))) # Set indentation width in twips
    table_indentation.set(qn('w:type'), 'dxa') # Set type attribute (dxa - twentieths of an inch)

    # Check if <w:tblInd> already exists, if so, replace it, otherwise append
    indent_element = table_properties.find(qn('w:tblInd'))
    if indent_element is not None:
        table_properties.replace(indent_element, table_indentation)
    else:
        table_properties.append(table_indentation)

def add_markdown_to_p(doc, paragraph, markdown):
    p = add_markdown_to_doc(doc, markdown) # add_markdown_to_doc now returns the paragraph
    paragraph.paragraph_format.keep_together = True # Keep lines in this paragraph together
    paragraph.paragraph_format.keep_with_next = True # Keep question text with next element

    for elem in doc.paragraphs[-1]._element:
        paragraph._element.append(elem)
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)


# only cell3 has markdown
def insert_item_start(doc, left_margin=Cm(1), cell1="", cell2="", md_cell3=""):

    table = doc.add_table(rows=0, cols=3)

    remove_table_padding(table)
    type_width = Cm(2.6)
    # table.autofit = False
    table.columns[0].width = left_margin
    table.columns[1].width = type_width
    table.columns[2].width = Cm(editable_width.cm - left_margin.cm - type_width.cm)

    row_cells = table.add_row().cells

    row_cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
    row_cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
    row_cells[2].paragraphs[0].paragraph_format.space_after = Pt(0)

    runner = row_cells[0].paragraphs[0].add_run(cell1)
    font = runner.font
    font.color.rgb = RGBColor(128, 128, 128)
    runner = row_cells[1].paragraphs[0].add_run(cell2)
    runner.bold = True

    add_markdown_to_doc(doc, md_cell3)
    point_text_para = row_cells[2].paragraphs[0]
    for elem in doc.paragraphs[-1]._element:
        point_text_para._element.append(elem)
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
    point_text_para.paragraph_format.keep_with_next = True # Keep question title with next element


def generate_docx_base64(test_data={}):
    """Generates a DOCX file in base64 format from test data."""

    doc = Document()
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN

    settings = test_data.get('settings', {})
    test_name = settings.get('test_name', 'Naamloze Toets')
    show_answers = settings.get('show_answers', True)


    doc.add_heading(test_name, level=1)
    doc.add_heading('Vragen', level=2)

    left_margin = Cm(1.3)
    question_number = 0

    for question in test_data.get('questions', []):
        question_number += 1

        # Horizontal rule before each question
        p = doc.add_paragraph()
        insertHR(p, left_margin)
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = True

        # Question context (if present)
        if question.get('question_context'):
            p_context = add_markdown_to_doc(doc, question['question_context'])
            p_context.paragraph_format.left_indent = left_margin
            p_context.paragraph_format.keep_with_next = True
            p_context.paragraph_format.keep_together = True

        # Question text (using the consistent table format)
        insert_item_start(doc, left_margin, cell1=f"{len(question.get('points', []))} pt", cell2=f"Vraag {question_number}", md_cell3=question["question_text"])


        # Answer and points (if show_answers is True)
        if show_answers:
            if question.get("answer_text"):
                p_answer_title = doc.add_paragraph()
                p_answer_title.paragraph_format.left_indent = left_margin
                p_answer_title.paragraph_format.keep_with_next = True
                p_answer_title.paragraph_format.keep_together = True
                run = p_answer_title.add_run("Correcte antwoord:")
                run.bold = True

                p_answer = add_markdown_to_doc(doc, question["answer_text"])
                p_answer.paragraph_format.left_indent = left_margin
                p_answer.paragraph_format.keep_with_next = True # Keep answer with points
                p_answer.paragraph_format.keep_together = True


            if question.get("points"):
                p_point_title = doc.add_paragraph()
                p_point_title.paragraph_format.left_indent = left_margin
                p_point_title.paragraph_format.keep_with_next = True
                p_point_title.paragraph_format.keep_together = True
                run = p_point_title.add_run("Punten:")
                run.bold = True

                table = doc.add_table(rows=0, cols=3)
                set_table_indentation(table, left_margin)
                table.columns[0].width = Cm(3)
                table.columns[1].width = Cm(3)
                table.columns[2].width = Cm(editable_width.cm - left_margin.cm - Cm(3).cm - Cm(4).cm)
                first_row = table.add_row()
                first_row.cells[0].paragraphs[0].add_run("Weging").bold = True
                first_row.cells[1].paragraphs[0].add_run("Punt").bold = True
                first_row.cells[2].paragraphs[0].add_run("Beschrijving").bold = True    
                
                
                # hdr_cells = table.rows[0].cells
                # hdr_cells[0].paragraphs[0].add_run('Weging').bold = True
                # hdr_cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
                # hdr_cells[0].paragraphs[0].paragraph_format.keep_with_next = True
                # hdr_cells[1].paragraphs[0].add_run('Punt').bold = True
                # hdr_cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
                # hdr_cells[1].paragraphs[0].paragraph_format.keep_with_next = True

                for point in question['points']:
                    row_cells = table.add_row().cells
                    # add_markdown_to_doc(doc, point.get('point_name', ''))
                    point_weight = point.get('point_weight', 1)
                    row_cells[0].paragraphs[0].text = str(point_weight) + ' pt'
                    row_cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
                    row_cells[0].paragraphs[0].paragraph_format.keep_with_next = True
                    
                    point_name = point.get('point_name', "")
                    row_cells[1].paragraphs[0].text = point_name
                    row_cells[1].paragraphs[0].paragraph_format.space_after = Pt(0)
                    row_cells[1].paragraphs[0].paragraph_format.keep_with_next = True
                    
                    add_markdown_to_p(doc, row_cells[2].paragraphs[0], point.get('point_text', ''))
                    row_cells[2].paragraphs[0].paragraph_format.space_after = Pt(0)
                    # Only keep with next if it's NOT the last point
                    row_cells[2].paragraphs[0].paragraph_format.keep_with_next = True if point != question['points'][-1] else False


    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    base64_output = base64.b64encode(buffer.read()).decode('utf-8')
    return base64_output