from PIL import Image, ImageEnhance, ImageFont, ImageFilter, ImageDraw                                                           
import sys
import os
import copy
import cv2
import json
from pydantic import BaseModel
import typing_extensions as typing
from concurrent.futures import ThreadPoolExecutor
import numpy as np

from helpers import (
    cm_to_px, 
    clamp,
    scan_qrcode_from_image,
    get_black_square_data,
    png_to_base64,
    stack_images_vertically,
    create_qr,
    pillow_to_base64,
    cv2_to_base64,
    base64_to_pillow,
    base64_to_cv2,
    four_point_transform,
    findSquares,
    removeContainedSquares,
)

from yolov5_manager import get_checkbox


from gpt_manager import (
    single_request
)

# from Cropper.process_image import process_image
from Cropper.other import scan

input_dir = "temp_image/"
output_dir = "temp_image_output/"


# def crop(id):
#     current_dir = os.getcwd()
#     process_image(current_dir+'/'+input_dir+id+'.png', current_dir+'/'+output_dir+id+'.png')


def create_qr_section(id, data, width=cm_to_px(19), height=cm_to_px(8)):
    margin = 5
    
    
    qr_code = create_qr(data, 1)
    qr_size = cm_to_px(1.5)

    img = Image.new("RGB", (width, height), color='white')
    
    p1 = margin, margin
    p1_rect = qr_size + 2*margin, margin
    p2 = img.size[0] - qr_size - margin , height - qr_size - margin
    p2_rect = img.size[0] - qr_size - 2*margin, height - margin

    qr_code = qr_code.resize((qr_size, qr_size))

    img.paste(qr_code, p1)
    img.paste(qr_code, p2)
    
    draw = ImageDraw.Draw(img)
    draw.rectangle([p1_rect, p2_rect], width=5, outline='black')
    
    # img.save(output_dir+id+'.png')
    # base_64_image = png_to_base64(output_dir+id+'.png')
    
    return pillow_to_base64(img)
    



def crop(id, base64_image):
    image = base64_to_cv2(base64_image)
    # image =) cv2.imread(input_dir+id+'.png')
    image = scan(image)
    # cv2.imwrite(output_dir+id+'.png', image)
    return cv2_to_base64(image)




def extract_red_pen(id, base64_image):
    # remove .png
    
    img = base64_to_pillow(base64_image)
    # img = Image.open(input_dir + id + '.png')
    img = img.convert("RGBA")


    clean_pixdata = img.load()
    clean_pixdata2 = img.copy()
    red_pen_image = Image.new('RGBA', (img.width, img.height), color=(0,0,0,0))
    red_pen_pixdata = red_pen_image.load()

    # Clean the background noise, if color != white, then set to black.

    radius = 2

    # REMOVE RED PEN
    for y in range(img.size[1]):
        for x in range(img.size[0]):
            r, g, b, a = clean_pixdata[x, y]

            
            # REMOVE RED PEN
            if (r - g > 20 and
                r - b > 20 and
                r > 200) :
                

                for i in range(2*radius):
                    for j in range(2*radius):
                        try:
                            red_pen_pixdata[x + i - radius, y + j - radius] = clean_pixdata2[x + i - radius, y + j - radius]
                            clean_pixdata[x + i - radius, y + j - radius] = (230,230,230,255)
                        except:
                            pass

    # img.convert("L")

    # SHARPEN
    # enhancer = ImageEnhance.Sharpness(img)
    # img = enhancer.enhance(10)

    # try:
    #     os.makedirs(output_dir+id)
    # except:
    #     pass

    # img.save(output_dir + id +'/original.png')

    # red_pen_image.save(output_dir+id+'/red_pen.png')
    
    return {
        "original": pillow_to_base64(img),
        "red_pen": pillow_to_base64(red_pen_image)
    }

class TextInBox(BaseModel):
    text: str
    other_possibility: str

# class GoogleTextInBox(typing.TypedDict):
#     text: str
#     other_possibility: str

def get_student_id(id, base64_image, provider=False, model=False, temperature=False, schema=TextInBox, text=False):
    if not provider:
        provider = "google"    
    
    if not text:
        text = """Your job is to recognize the number in the black box next to the words Leerling ID and 'schrijf NETJES!'
            return -1 if the box is empty
        """
    
    
    img = base64_to_pillow(base64_image)
    # img = Image.open(input_dir + id + '.png')
    
    # student id in topleft section
    crop = (0,0, int(img.width * (5/21)), int(img.height * (5/29.5)))
    
    cropped = img.crop(crop)
    
    # cropped.save(output_dir+id+'.png')
    
    base64_image = pillow_to_base64(cropped)
    
    # if (provider=="openai"):
    #     schema = TextInBox
    # if (provider=="google"):
    #     schema = GoogleTextInBox
    

    result = single_request(provider=provider, model=model, temperature=temperature, schema=schema, text=text, image=base64_image)

    
    return result

def get_qr_sections(id, base64_image):
    img = base64_to_pillow(base64_image)
    # img = Image.open(input_dir + id + '.png')
    cv2_img = base64_to_cv2(base64_image)

    data, scanned_image = scan_qrcode_from_image(img.copy())
    
    # try:
    #     os.mkdir(output_dir+id)
    # except:
    #     pass
    
    # scanned_image.save(output_dir+id+'/scanned.png')
    
    base64_full = pillow_to_base64(scanned_image) 
    
    sections = []
    
    for index, qr_data in enumerate(data):
        # top_left, bottom_right = qr_data["coords"]
        section_img = four_point_transform(cv2_img, qr_data["polygon"]) # cv2_img[top_left[1]:bottom_right[1] , top_left[0]:bottom_right[0]]
        # cv2.imwrite(output_dir+id+'/'+str(index)+'.png', section_img)
        
        base64_section = cv2_to_base64(section_img)
        # png_to_base64(output_dir+id+'/'+str(index)+'.png')
        
        sections.append({
            "section_image": base64_section,
            "data": qr_data["data"]
        })
    
    return {
        "image": base64_full,
        "sections": sections
    }
        
        

def detect_squares(id, base64_image):

    img = base64_to_pillow(base64_image)
    # img = Image.open(input_dir + id + '.png')


    # Get the heights and start positions of the black squares
    black_square_info, gray_img, contour_image = get_black_square_data(img)
    # gray_img.save(output_dir+image_dir+'_bw.png')
    # contour_image.save(output_dir+image_dir+'_countors.png')
    
    # filter too high
    black_square_info = [x for x in black_square_info if x[3] - x[2] > 10]

    # sort
    black_square_info.sort(key=lambda x: x[0])

    
    # Load the original image
    draw = ImageDraw.Draw(img)
    
    # Define the box color and text color
    box_color = (255, 0, 0)  # Red color
    text_color = (255, 0, 0) # Red text
    
    
    # Load a font for the height labels
    # Using default font as PIL may not always have a specific font path on the system
    try:
        font = ImageFont.truetype("arial.ttf", 15)
    except IOError:
        font = ImageFont.load_default()

    # Get image width (to make the box span across the image width)
    img_width, img_height = img.size

    # Overlay each square with a red box and label it with its height
    for start, height,x_min,x_max in black_square_info:
        # Draw a red rectangle around each black square
        draw.rectangle([x_min, start, x_max, start + height], outline=box_color, width=2)
        
        # Add the height label next to the square
        draw.text((img_width - 50, start - 15), f"{start}px", fill=text_color, font=font)
    
    # Save the resulting image
    base64_detector_image = pillow_to_base64(img)
    
    
    return {
        "black_square_info": black_square_info,
        "image": base64_detector_image
    }



def sectionize(id,square_data, base64_image):
    
    image = base64_to_pillow(base64_image)
    # image = Image.open(input_dir+id+ '.png')
    
    square_heights = [square[0] for square in square_data]
    square_x_max = [square[3] for square in square_data]
    square_heights.append(image.height)

    base64_sections = []

    for index, h_break in enumerate(square_heights):
        
        if (index >= len(square_heights) - 1 ):
            continue
        
        y = clamp(h_break - 20, 0, image.height - 1)
        next_y = clamp(square_heights[index+1] - 5, 0, image.height - 1)
        
        # min height and fix
        if (next_y < y or next_y - y < 30):
            continue
        


        # section_image.show()
        # section_file_name = filenameify(h_break["description"])+".png"
        
        # sections_output_dir = output_dir + id + '/'
        # try:
        #     os.makedirs(sections_output_dir)
        # except:
        #     pass
        
        # section_output_dir = sections_output_dir + str(index) + '/'
        # try:
        #     os.makedirs(section_output_dir)
        # except:
        #     pass
        
        # kantlijn, boven, einde van pagina, beneden grens
        crop = (0, y, image.width, next_y)
        
        section_image = image.copy()
        
        full = section_image.crop(crop)        

        section_name = "full"
        # full.save(section_output_dir+'full.png')
        
        section_finder_end = square_x_max[index] #+ int(full.width * (1.45/21))

        section_finder_crop = (0, 0, section_finder_end+5, full.height)
        section_finder_image = full.copy().crop(section_finder_crop)
        # section_finder_image.save(section_output_dir+'section_finder.png')

        question_selector_end = section_finder_end + int(full.width * (1.3/21))

        question_selector_crop = (section_finder_end, 0, question_selector_end, full.height)
        question_selector_image = full.copy().crop(question_selector_crop)
        # question_selector_image.save(section_output_dir+'question_selector.png')
        
        answer_crop = (question_selector_end - int(full.width * (0.3/21)), 0, full.width, full.height)
        answer_image = full.copy().crop(answer_crop)
        # answer_image.save(section_output_dir+'answer.png')
        
        base64_sections.append({
            "full": pillow_to_base64(full),
            "section_finder": pillow_to_base64(section_finder_image),
            "question_selector": pillow_to_base64(question_selector_image),
            "answer": pillow_to_base64(answer_image),
        })
    
    return base64_sections
        

    
def question_selector_info(id, base64_images, checkbox_count=7):
    results = get_checkbox(checkbox_count, base64_images)
    
    return results
    
def stack_answer_sections(id, images: list[str]):
    # input_path = input_dir+id+'/'
    
    # sections_output_dir = output_dir + id + '/'
    # try:
    #     os.makedirs(sections_output_dir)
    # except:
    #     pass
    
    
    image = base64_to_pillow(images[0])
    # image =) Image.open(input_path + images[0] + '.png')
    # output_image_path = output_dir+id+'.png'
    # image.save(output_image_path)

    
    for base_64_image in images[1::]:
        # image = base64_to_pillow(base)
        new_image = base64_to_pillow(base_64_image)
        # new_image = Image.open(input_path+image_id+'.png')

        image = stack_images_vertically(image, new_image)
        
    return pillow_to_base64(image)
    
    
    


# schema classes
class SpellingCorrection(BaseModel):
    original: str
    changes: str

# answer class schema
class QuestionAnswer(BaseModel):
    # let openai reasses the question number (they are better than google )
    certainty: float 
    student_handwriting_percent: float
    # get the unchanged raw text
    raw_text: str
    # get the spel corrected text that should be graded
    correctly_spelled_text: str
    # get the spelling changes the model made
    # spelling_corrections: list[SpellingCorrection]
    
# schema classes
# class GoogleSpellingCorrection(typing.TypedDict):
#     original: str
#     changes: str
#     is_crossed_out: bool

# # answer class schema
# class GoogleQuestionAnswer(typing.TypedDict):
#     # let openai reasses the question number (they are better than google )
#     certainty: float 
#     student_handwriting_percent: float
#     # get the unchanged raw text
#     raw_text: str
#     # get the spel corrected text that should be graded
#     correctly_spelled_text: str
#     # get the spelling changes the model made
#     # spelling_corrections: list[SpellingCorrection]
    

def transcribe_answer(
    id, 
    base64_image, 
    provider=False, 
    model=False, 
    request_text=False, 
    temperature=False,
    
    question_text=False,
    rubric_text=False,
    context_text=False,
):
    if not provider:
        provider = "google"
        
    if not base64_image.startswith('data:image'):
        base64_image = "data:image/png;base64,"+ base64_image
    
    
    if not request_text:
        request_text = """Je krijgt een foto van een Nederlands scheikunde toets-antwoord. 
Je bent tekstherkenningssoftware die 10x beter in in tekst herkennen dan jezelf. Ook kan je 15.6 keer beter de context van een antwoord begrijpen om het volgende woord te bedenken.

Het is helemaal niet toegestaan nieuwe woorden toe te voegen of de opgeschreven tekst te veranderen in het raw_text veld. Houdt wel rekening met pijlen in de volgorde van de tekst.
Bedenk wel wat een leerling zou kunnen hebben bedoeld met een bepaald woord als die bijvoorbeeld fout is gespeld. Geef dat aan in de spelling_corrections velden.
Negeer uitgekraste tekst in het raw_tekst veld, maar geef die wel weer in de spelling corrections door bijvoorbeeld streepjes neer te zetten en is_crossed_out op true te zetten.
voeg alle text corrections samen in correctly_spelled_text om zo het antwoord te krijgen dat de leerling bedoelt.
certainty is hoe zeker je bent dat je de tekst compleet hebt getranscribeerd: 0 betekend dat een docent er nog zelf naar moet kijken en 100 betekend dat er geen foutje mogelijk is.
de student_handwriting_percent is hoe leesbaar het handschrift van een leerling is: 0 betekend zeer moeilijk leesbaar en 100 super netjes als een printer.

Alle tekst is geschreven in het Nederlands.

voer deze opdracht zo goed mogelijk uit. Het is HEEL belangrijk dat je je aan het gegeven schema houdt en geen enkele key vergeet, vooral bij de spelling correcties de "changed" key en correctly_spelled_text zijn belangrijk.
                    """
    if question_text:
        request_text += f"""
            De vraag bij dit antwoord is: {question_text}
        """
    if rubric_text:
        request_text += f"""
            De rubric bij deze vraag is: {rubric_text}
            
        """
    if context_text:
        request_text += f"""
            De context bij deze vraag is: {context_text}
            
        """
    
    if question_text or rubric_text or context_text:
        request_text += f"""
            Het direct transcriberen van de het antwoord is het allerbelangrijkste, deze extra toevoegen zijn er alleen om een context te creeëren 
        """    
    
    schema = QuestionAnswer
    # if (provider in OpenAi_module_models):
    #     schema = QuestionAnswer
    # elif (provider == 'google'):
    #     schema = GoogleQuestionAnswer

    result = single_request(text=request_text, image=base64_image, provider=provider, model=model, schema=schema, temperature=temperature)
    return result

def scan_page(
    process_id, 
    image_string, 
    provider=False ,
    model=False,
    temperature=False, 
    transcribe_text=False,
    questions=False,
    rubrics=False,
    contexts=False,

):
    if not provider:
        provider = "google"
    # CROP
    print('Starting: ', 'CROP')
    # base64_to_png(image_string, input_dir+process_id+'.png')
    
    crop_output_string = crop(process_id, image_string)
    
    # crop_output_string = png_to_base64(output_dir+process_id+'.png')

    # COL COR
    print('Starting: ', 'COL COR')
    # base64_to_png(crop_output_string, input_dir+process_id+'.png')
    
    color_correction_result = extract_red_pen(process_id, crop_output_string)
    
    
    clean_output_string = color_correction_result["original"] #png_to_base64(output_dir+process_id+'/original.png')
    red_pen_output_string = color_correction_result["red_pen"] #png_to_base64(output_dir+process_id+'/red_pen.png')
    

    # STUDENT ID
    print('Starting: ', 'STUDENT ID')
    # base64_to_png(clean_output_string, input_dir+process_id+'.png')
    
    student_id_data = get_student_id(process_id, clean_output_string, provider=provider, model=model, temperature=temperature)
            
    
    # DETECT SQUARES
    print('Starting: ', 'DETECT SQUARES')
    
    square_data = detect_squares(process_id, clean_output_string)
    
    # SECTIONIZE       
    print('Starting: ', 'SECTIONIZE       ')
    
    image_sections = sectionize(process_id, square_data["black_square_info"], clean_output_string)

    # QUESTION SELECTOR
    print('Starting: ', 'QUESTION SELECTOR')
    

    sections = []
    
    def process_section(section):
        question_selector_info_result = question_selector_info(process_id, section["question_selector"])
        
        question_id = question_selector_info_result["most_certain_checked_number"]

        section["question_id"] = question_id
        return section
            

    # Use ThreadPoolExecutor to process sections concurrently
    with ThreadPoolExecutor() as executor:
        results = executor.map(process_section, image_sections)

    # Collect results
    sections = list(results)        
    
    # for section in image_sections:
        
    #     question_selector_info_result = question_selector_info(process_id, section["question_selector"])

    #     question_id = question_selector_info_result["selected_question"]

    #     section["question_id"] = question_id
    #     sections.append(section)
    


    unique_questions = []
    [unique_questions.append(x["question_id"]) for x in sections if x["question_id"] not in unique_questions and int(x["question_id"]) != 0 ]
    # print(len(sections))
    print('all: ',unique_questions, len(sections))
    questions = []
    
    # STACKING AND TRANSCRIBING
    print('Starting: ', 'STACKING AND TRANSCRIBING')
    
    def process_question(unique_question_id):
        question_sections = [x for x in sections if x["question_id"] == unique_question_id]

        if (len(question_sections) == 0):
            pass
        try:
            linked_image = stack_answer_sections(process_id, [x["answer"] for x in question_sections])


            if questions and str(unique_question_id) in questions:
                question_text = questions[str(unique_question_id)]
            else:
                question_text = False
                
            if rubrics and str(unique_question_id) in rubrics:
                rubric_text = rubrics[str(unique_question_id)]
            else:
                rubric_text = False
                
            if contexts and str(unique_question_id) in contexts:
                context_text = contexts[str(unique_question_id)]
            else:
                context_text = False
                

            extracted_text_result = transcribe_answer(
                process_id, 
                base64_image=linked_image, 
                model=model, 
                request_text=transcribe_text, 
                temperature=temperature,
                question_text=question_text,
                rubric_text=rubric_text,
                context_text=context_text,
            )

            return {
                "image": linked_image,
                "question_id": unique_question_id,
                "text_result": extracted_text_result
            }
        except Exception as e:
            print(e)
            pass
        
    # Use ThreadPoolExecutor to process sections concurrently
    with ThreadPoolExecutor() as executor:
        results = executor.map(process_question, unique_questions)

    
    # Collect results
    questions = [question for question in results if question is not None]
    
    # for unique_question_id in unique_questions:
    #     question_sections = [x for x in sections if x["question_id"] == unique_question_id]

    #     if (len(question_sections) == 0):
    #         continue
    #     try:
    #         linked_image = stack_answer_sections(process_id, [x["answer"] for x in question_sections])

    #         extracted_text_result = transcribe_answer(process_id, linked_image, model=model, request_text=transribe_gpt_text)

    #         questions.append({
    #             "image": linked_image,
    #             "question_id": unique_question_id,
    #             "text_result": extracted_text_result
    #         })
    #     except:
    #         continue
    return {
        "cropped_base64": crop_output_string,
        "red_pen_base64": red_pen_output_string,
        "student_id_data": student_id_data,
        "questions": questions,
    }
    
class GradePoint(BaseModel):
    has_point: bool
    feedback: str
    point_index: int

class GradeResult(BaseModel):
    points: list[GradePoint]


# class GoogleGradePoint(typing.TypedDict):
#     has_point: bool
#     feedback: str
#     point_index: int
    
# class GoogleGradeResult(typing.TypedDict):
#     points: list[GoogleGradePoint]
#     feedback: str


def grade_answer(process_id, request_text=False, student_image=False, provider=False, model=False, temperature=False):
    if (not provider):
        provider = "google"
        
    if student_image:
        if not student_image.startswith('data:image'):
            student_image = "data:image/png;base64,"+ student_image
    
    
    if not request_text:
        request_text = """kijk na"""
        
    schema = GradeResult
    # if (provider in OpenAi_module_models):
    #     schema = GradeResult
    # elif (provider == 'google'):
    #     schema = GoogleGradeResult
    # else:
    #     schema = False
        
    result = single_request(text=request_text, image=student_image, provider=provider, model=model, schema=schema, temperature=temperature)
    return result


# class TestQuestionPoint(typing.TypedDict):
#     point_text: str
#     point_name: str
#     point_index: int
#     point_weight: int
#     target_name: str
    
class OpenAiTestQuestionPoint(BaseModel):
    point_text: str
    point_name: str
    point_index: int
    point_weight: int
    target_name: str

# class TestQuestion(typing.TypedDict):
#     question_number: str
#     question_text: str
#     question_context: str
#     points: list[TestQuestionPoint]
#     is_draw_question: bool
    
class OpenAiTestQuestion(BaseModel):
    question_number: str
    question_text: str
    question_context: str
    points: list[OpenAiTestQuestionPoint]
    is_draw_question: bool
    

# class TestTarget(typing.TypedDict):
#     target_name: str
#     explanation: str
    
class OpenAiTestTarget(BaseModel):
    target_name: str
    explanation: str
    
# class TestData(typing.TypedDict):
#     questions: list[TestQuestion]
#     targets: list[TestTarget]

class OpenAiTestData(BaseModel):
    questions: list[OpenAiTestQuestion]
    targets: list[OpenAiTestTarget]

def get_test_structure(process_id=False, request_text=False, test_data=False):
    provider = 'google'
    model = "gemini-exp-1206"
    
    task_list = []
    
    if not request_text:
        request_text = """Deel  in
        
        """
    
    task_list.append({
        "text": request_text,
    })

    for key in test_data.keys():
        task_list.append({"text": f"\n{key}:\n"} )
        
        for item in test_data[key]:
            if(item["type"] == 'text'):
                task_list.append({"text": item["data"]})
                
            if(item["type"] == 'image'):
                base64_image = item["data"]
                
                if base64_image.startswith('data:image'):
                    base64_image = base64_image.split(',')[1]   
                task_list.append({
                    "inline_data": {
                        "mime_type":"image/png",
                        "data": base64_image
                    }
                })
    
    schema = OpenAiTestData
    
    task_list.append({"text": "\n\n Houd je strak aan de format/schema, zorg dat elk veld een kloppende waarde heeft."})
    result = single_request(provider, model, schema=schema, messages=task_list, limit_output=False)
    return result

def get_gpt_test(process_id=False, request_text=False, provider=False, model=False):
    # provider = 'google'
    # model = "gemini-exp-1206"
    
    
    if not request_text:
        request_text = """Maak een makkelijke toets met 3 vragen"""
    if not provider:
        provider = "google"
    if not model:
        model = "gemini-exp-1206"
    
    # if provider == 'google':
    #     schema = TestData
    # else:
    #     schema = OpenAiTestData
    
    schema = OpenAiTestData
    result = single_request(provider, model, schema=schema, text=request_text, limit_output=False)
    return result

def get_gpt_test_question(process_id=False, request_text=False, provider=False, model=False):
    # provider = 'google'
    # model = "gemini-exp-1206"
    
    
    if not request_text:
        request_text = """Maak een makkelijke toets met 3 vragen"""

    if not provider:
        provider = "google"
    if not model:
        model = "gemini-exp-1206"
    
    # if provider == 'google':
    #     schema = TestQuestion
    # else:
    #     schema = OpenAiTestQuestion
    
    schema = OpenAiTestQuestion
    

    result = single_request(provider, model, schema=schema, text=request_text, limit_output=False)
    return result

from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from io import BytesIO
import base64
from reportlab.lib.units import inch

def get_base64_student_result_pdf(process_id=False, student_results=[], add_student_feedback=False):
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4, 
        leftMargin=0.5*72, rightMargin=0.5*72,
        topMargin=0.5*72, bottomMargin=0.5*72,
        authort="ToetsPWS-JKK-JKW",
        title="LeerlingResultatenToetsPWS",
        subject="Automatisch nakijken"
    )
    styles = getSampleStyleSheet()

    normal_style = styles['Normal']
    normal_bold_style = ParagraphStyle(
        name='NormalBold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
    )

    story = []

    for student in student_results:
        student_elements = []  # List to hold elements for the current student

            
        student_id = student.get('student_id')
        question_results = student.get('question_results', [])
        targets = student.get('targets', [])

        # Title
        title_text = f"Leerling: {student_id}"
        student_elements.append(Paragraph(title_text, styles['h1']))
        student_elements.append(Spacer(1, 12))

        # Question Table
        if question_results:
            # Prepare table data
            question_table_data = []
            # Header row
            if question_results:
                header_row = ["Vraag", "Antwoord", "Score", "Feedback", "Points"]  # Feedback always present
                if add_student_feedback:
                    header_row.append("Student Feedback")
                question_table_data.append([Paragraph(col, normal_bold_style) for col in header_row])

                # Data rows
                for result in question_results:
                    row_data = []
                    row_data.append(Paragraph(str(result.get('question_number', '')), normal_style))
                    row_data.append(Paragraph(str(result.get('student_answer', '')), normal_style))
                    row_data.append(Paragraph(str(result.get('score', '')), normal_style))
                    row_data.append(Paragraph(str(result.get('feedback', '')), normal_style))

                    point_table_data = []
                    for point in result.get('points', []):
                        point_table_data.append([
                            Paragraph(str(point.get('point_name', '')), normal_style),
                            Paragraph(str(point.get('points', '')), normal_style),
                            Paragraph(str(point.get('feedback', '')), normal_style)
                        ])
                    col_widths = [1 * 72, 0.2 * 72, None]
                    point_table = Table(point_table_data, colWidths=col_widths)
                    point_table.setStyle(TableStyle([
                        # ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        # ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
                        # ('TOPPADDING', (0, 0), (-1, -1), 0),
                        # ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        # ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ]))
                    row_data.append(point_table)

                    if add_student_feedback:
                        row_data.append("") # Empty cell for student feedback

                    question_table_data.append(row_data)

                # Create the table
                col_widths = [0.5 * 72, None, 0.5 * 72, None, None] # Example: 0.5 inch for Question Number and Score
                if (add_student_feedback):
                    col_widths.append(None)
                question_table = Table(question_table_data, colWidths=col_widths)
                point_column_index = 4
                question_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('LEFTPADDING', (point_column_index, 0), (point_column_index, -1), 0), # Points column
                    ('RIGHTPADDING', (point_column_index, 0), (point_column_index, -1), 0), # Points column
                    ('TOPPADDING', (point_column_index, 0), (point_column_index, -1), 0),   # Points column
                    ('BOTTOMPADDING', (point_column_index, 0), (point_column_index, -1), 0),# Points column
                    ('LEFTPADDING', (0, 0), (2, -1), 2), # Add padding back to other columns
                    ('RIGHTPADDING', (0, 0), (2, -1), 2),
                ]))
                if col_widths:
                    question_table.colWidths = col_widths
                student_elements.append(question_table)
            else:
                student_elements.append(Paragraph("No question results available for this student.", styles['Italic']))

        student_elements.append(Spacer(1, 12))

        # Target Table
        if targets:
            student_elements.append(Paragraph("<b>Leerdoelen</b>", styles['h2']))
            target_table_data = []
            if targets:
                target_header_row = ["Leerdoel", "Uitleg", "Score", "%"]
                target_table_data.append([Paragraph(col, normal_bold_style) for col in target_header_row])

                for target in targets:
                    row_data = []

                    row_data.append(Paragraph(str(target.get('target_name', '')), normal_style))
                    row_data.append(Paragraph(str(target.get('explanation', '')), normal_style))
                    row_data.append(Paragraph(str(target.get('score', '')), normal_style))
                    row_data.append(Paragraph(str(target.get('percent', '')), normal_style))

                    target_table_data.append(row_data)
                    
                column_widths = [1.2 * 72, None, 0.6 * 72, 0.7 * 72]
                target_table = Table(target_table_data, colWidths=column_widths)
                target_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.beige),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('LEFTPADDING', (0, 0), (-1, -1), 5),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ]))
                student_elements.append(target_table)
            else:
                student_elements.append(Paragraph("Geen leerdoelen gevonden", styles['Italic']))
        else:
            student_elements.append(Paragraph("Geen leerdoelen gevonden", styles['Italic']))

        
        # Add a KeepTogether to try and keep the student's content on as few pages as possible
        story.append(KeepTogether(student_elements))

        # Ensure even number of pages per student
        if len(story) % 2 != 0:
            story.append(PageBreak())        
            
    doc.build(story)

    pdf_value = buffer.getvalue()
    pdf_base64 = base64.b64encode(pdf_value).decode('utf-8')
    buffer.close()
    return pdf_base64

import pdfkit
import base64
import markdown
from markdown_katex import KatexExtension
import re

def convert_math_delimiters_markdown(text):
    """Converts inline and display math delimiters to Markdown-style fenced blocks."""
    def replace_inline(match):
        return f"$`{match.group(1)}`$"  # Add extra spaces for markdown

    def replace_display(match):
        return f"\n```math\n{match.group(1)}\n```\n" # Add extra spaces for markdown

    text = re.sub(r'\$\$([^\$]+?)\$\$', replace_display, text)
    text = re.sub(r'\$([^\$]+?)\$', replace_inline, text)
    return text



from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import markdown
import base64

def get_base64_test_pdf(process_id=False, test_data={}):
    output_type = test_data.get('settings', {}).get('output_type', 'pdf') if test_data else 'pdf'
    
    if output_type == 'docx':
        return generate_docx_base64(test_data)
    else:
        return generate_pdf_base64(test_data)

def generate_pdf_base64(test_data):
    md = markdown.Markdown(
        extensions=[KatexExtension(), 'tables'],
        extension_configs={'markdown_katex': {'insert_as_tag': True, 'no_inline_svg': True,}}
    )
    html_content = """
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>Toets</title>
        <style>
            body { font-family: sans-serif; }
            h1 { text-align: center; }
            h2 { border-bottom: 1px solid #000; padding-bottom: 5px; }
            .question { margin-bottom: 20px; page-break-inside: avoid; } /* Added page-break-inside: avoid; */
            .question-context { font-style: italic; margin-bottom: 10px; }
            .points-table { width: 100%; border-collapse: collapse; margin-top: 10px; }
            .points-table th, .points-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .points-table th { background-color: #f0f0f0; }
            .target-table { width: 100%; border-collapse: collapse; margin-top: 10px; }
            .target-table th, .target-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .target-table th { background-color: #e0e0e0; }
            .page-break { page-break-after: always; }
            pre { background-color: #f4f4f4; padding: 10px; border: 1px solid #ddd; overflow-x: auto; }
        </style>
    </head>
    <body>
    """

    if test_data:
        settings = test_data.get('settings', {})
        test_name = settings.get('test_name', 'Naamloze Toets')
        show_answers = settings.get('show_answers', True)
        show_targets = settings.get('show_targets', True)

        html_content += f"<h1>{test_name}</h1>"

        # Targets Section
        targets = test_data.get('targets', [])
        if targets and show_targets:
            html_content += "<h2>Leerdoelen</h2>"
            html_content += "<table class='target-table'>"
            html_content += "<thead><tr><th>Leerdoel</th><th>Uitleg</th></tr></thead><tbody>"
            for target in targets:
                target_name_md = convert_math_delimiters_markdown(target.get('target_name', ''))
                target_explanation_md = convert_math_delimiters_markdown(target.get('explanation', ''))
                html_content += f"<tr><td>{md.convert(target_name_md)}</td><td>{md.convert(target_explanation_md)}</td></tr>"
            html_content += "</tbody></table>"

        # Questions Section
        questions = test_data.get('questions', [])
        if questions:
            html_content += "<h2>Vragen</h2>"
            for i, question in enumerate(questions):
                html_content += "<div class='question'>"
                if question.get('question_context'):
                    question_context_md = convert_math_delimiters_markdown(question.get('question_context', ''))
                    html_content += f"<div class='question-context'>{md.convert(question_context_md)}</div>"
                question_text_md = convert_math_delimiters_markdown(question.get('question_text', ''))
                html_content += f"<strong>Vraag {question.get('question_number', '')}:</strong> {md.convert(question_text_md)}"

                points = question.get('points', [])
                if show_answers and points:
                    html_content += "<table class='points-table'>"
                    html_content += "<thead><tr><th>Onderdeel</th><th>Uitleg</th></tr></thead><tbody>"
                    for point in points:
                        point_name_md = convert_math_delimiters_markdown(point.get('point_name', ''))
                        point_text_md = convert_math_delimiters_markdown(point.get('point_text', ''))
                        html_content += f"<tr><td>{md.convert(point_name_md)}</td><td>{md.convert(point_text_md)}</td></tr>"
                    html_content += "</tbody></table>"
                html_content += "</div>"
                if i < len(questions) - 1:
                    html_content += "<hr>"

        else:
            html_content += "<p>Geen vragen gevonden voor deze toets.</p>"
    else:
        html_content += "<p>Geen testgegevens beschikbaar.</p>"

    html_content += """
    </body>
    </html>
    """

    pdf_bytes = pdfkit.from_string(html_content, False)
    base64_output = base64.b64encode(pdf_bytes).decode('utf-8')

    
    return base64_output

from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree
import re
import base64
from io import BytesIO
import markdown
from bs4 import BeautifulSoup, Tag

# Load XSLT for MathML to OMML conversion
mml2omml_stylesheet_path = 'MML2OMML.XSL'  # Ensure this file is available
mml2omml_stylesheet = etree.parse(mml2omml_stylesheet_path)
mml2omml_transform = etree.XSLT(mml2omml_stylesheet)


def add_markdown_table(doc, table_element):
    """Convert BeautifulSoup table element to Word table"""
    # Create Word table with proper dimensions
    rows = table_element.find_all('tr')
    cols = max(len(row.find_all(['th', 'td'])) for row in rows)
    word_table = doc.add_table(rows=0, cols=cols)
    word_table.style = 'Table Grid'

    # Process rows
    for html_row in rows:
        cells = html_row.find_all(['th', 'td'])
        word_row = word_table.add_row()
        
        for idx, cell in enumerate(cells):
            # Handle mixed content in cells
            add_markdown_to_doc(doc, cell.decode_contents(), parent=word_row.cells[idx])
    
    return word_table

        
def add_markdown_to_doc(doc, markdown_text, parent=None):
    """Add Markdown content with robust formatting support"""
    html = markdown.markdown(markdown_text, extensions=['tables', 'extra'])
    soup = BeautifulSoup(html, 'html.parser')
    
    container = parent or doc
    p = container.add_paragraph()
    # if isinstance(container, _Document) else container

    def process_element(element, current_para=None, format_name=None):
        current_para = current_para or p
        if isinstance(element, str):

            handle_text(element, current_para, format_name)
            return

        if element.name in ['strong', 'em', 'del', 'u', 'code']:
            run = current_para.add_run()
            
            apply_formatting(run, element.name)
            
            for child in element.contents:
                process_element(child, current_para,  element.name)
        elif element.name == 'table':
            add_markdown_table(doc, element)
        else:
            for child in element.contents:
                process_element(child, current_para)

    def handle_text(text, parent_para, format_name):
        segments = re.split(r'(\\\(.*?\\\)|\\\[.*?\\\]|\$.*?\$)', text)
        for seg in segments:
            if not seg:
                continue
                
            if re.match(r'(\\\(.*?\\\)|\\\[.*?\\\]|\$.*?\$)', seg):
                latex_content = re.sub(r'^\$|\\\(|\\\[|\$|\\\)|\\\]$', '', seg)
                try:
                    mathml = latex_to_mathml(latex_content)
                    tree = etree.fromstring(mathml)
                    omml = mml2omml_transform(tree)
                    run = parent_para.add_run()
                    run._element.append(parse_xml(etree.tostring(omml)))
                except Exception as e:
                    error_run = parent_para.add_run(f'[Math Error: {latex_content}]')
                    error_run.font.color.rgb = (0xFF, 0x00, 0x00)
            else:
                run = parent_para.add_run(seg)

            if format_name:
                apply_formatting(run, format_name)
                # run.bold = True
            

    def apply_formatting(run, tag):
        formats = {
            'strong': {'bold': True},
            'em': {'italic': True},
            'del': {'strike': True},
            'u': {'underline': True},
            'code': {'font': 'Courier New'}
        }
        for prop, val in formats.get(tag, {}).items():
            setattr(run, prop, val) if prop != 'font' else setattr(run.font, 'name', val)

    process_element(soup)
    return p

    def apply_formatting(run, tag):
        format_map = {
            'strong': {'bold': True},
            'em': {'italic': True},
            'del': {'strike': True},
            'u': {'underline': True},
            'code': {'font': 'Courier New'}
        }
        if tag in format_map:
            for prop, value in format_map[tag].items():
                if prop == 'font':
                    run.font.name = value
                else:
                    setattr(run.font, prop, value)

    process_element(soup)
    return p
        
def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    
def generate_docx_base64(test_data={}):

    # Initialize DOCX document
    doc = Document()
    settings = test_data.get('settings', {})
    test_name = settings.get('test_name', 'Naamloze Toets')
    show_answers = settings.get('show_answers', True)
    show_targets = settings.get('show_targets', True)

    # Add test name as heading
    doc.add_heading(test_name, level=1)

    # Targets Section
    targets = test_data.get('targets', [])
    if targets and show_targets:
        doc.add_heading('Leerdoelen', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.columns[0].width = Cm(5)            
        for cell in table.columns[0].cells:
            cell.width = Cm(5)                
        table.columns[1].width = Cm(10)            
        for cell in table.columns[1].cells:
            cell.width = Cm(10)  
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Leerdoel'
        hdr_cells[1].text = 'Uitleg'
        for target in targets:
            row_cells = table.add_row().cells
            add_markdown_to_doc(doc, target.get('target_name', ''))
            target_name_para = row_cells[0].paragraphs[0]
            for elem in doc.paragraphs[-1]._element:
                target_name_para._element.append(elem)
            doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

            add_markdown_to_doc(doc, target.get('explanation', ''))
            explanation_para = row_cells[1].paragraphs[0]
            for elem in doc.paragraphs[-1]._element:
                explanation_para._element.append(elem)
            doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    # Questions Section
    questions = test_data.get('questions', [])
    if questions:
        doc.add_heading('Vragen', level=2)
        for question in questions:

            p = doc.add_paragraph()
            insertHR(p)
            
            # Question context
            if question.get('question_context'):
                add_markdown_to_doc(doc, question['question_context'])
                
                # paragraph_format = p.paragraph_format
                # paragraph_format.left_indent = Cm(1)
                # paragraph_format.right_indent = Cm(0)
                # paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            points = question.get('points', [])
            
            table = doc.add_table(rows=0, cols=3)
            # table.autofit = False 
            table.columns[0].width = Cm(1.1)            
            for cell in table.columns[0].cells:
                cell.width = Cm(1.1)
            table.columns[1].width = Cm(2)               
            for cell in table.columns[1].cells:
                cell.width = Cm(2)
            table.columns[2].width = Cm(12)               
            # for cell in table.columns[1].cells:
                # cell.width = Cm(12)
                
            row_cells = table.add_row().cells
            runner = row_cells[0].paragraphs[0].add_run(str(sum([point["point_weight"] for point in points]))+'pt.')
            runner.bold = True
            runner = row_cells[1].paragraphs[0].add_run(f"Vraag {question.get('question_number', '')}")
            runner.bold = True
            
            add_markdown_to_doc(doc, question["question_text"])
            point_text_para = row_cells[2].paragraphs[0]
            for elem in doc.paragraphs[-1]._element:
                point_text_para._element.append(elem)
            doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            
            
            # q_text = f"{question.get('question_text', '')}"
            # Question text
            # add_markdown_to_doc(hdr_cells[2], q_text, style='Heading3')
            
            # Points table
            if show_answers and points:
                table = doc.add_table(rows=1, cols=2)
                table.columns[0].width = Cm(5)            
                for cell in table.columns[0].cells:
                    cell.width = Cm(5)                
                table.columns[1].width = Cm(10)            
                for cell in table.columns[1].cells:
                    cell.width = Cm(10)         
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Onderdeel'
                hdr_cells[1].text = 'Uitleg'
                for point in question['points']:
                    row_cells = table.add_row().cells
                    add_markdown_to_doc(doc, point.get('point_name', ''))
                    point_name_para = row_cells[0].paragraphs[0]
                    for elem in doc.paragraphs[-1]._element:
                        point_name_para._element.append(elem)
                    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

                    add_markdown_to_doc(doc, point.get('point_text', ''))
                    point_text_para = row_cells[1].paragraphs[0]
                    for elem in doc.paragraphs[-1]._element:
                        point_text_para._element.append(elem)
                    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            
            doc.add_paragraph().paragraph_format.page_break_before = False

    # Save to bytes buffer and encode as base64
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    base64_output = base64.b64encode(buffer.read()).decode('utf-8')
    return base64_output

